#!/usr/bin/env python3
"""
docx_to_md.py — Konverterar ett TKB Word-dokument till sektionerade Markdown-filer.

Användning:
    python3 docx_to_md.py <docx-fil> <output-katalog>

Output-katalogen fylls med:
    sections/
        1-inledning.md
        2-versionsinformation.md
        ...
        7-tjanstekontrakt.md
        full-document.md     ← hela dokumentet som en fil
    images/
        img_001.png
        img_002.png
        ...
    structure.json           ← rubrikhierarki för navigation

Bildfilerna kopieras till output-katalogen/images/ och refereras med
relativa sökvägar (images/img_001.png) i markdown-texten.

Tabeller konverteras till Markdown-tabellsyntax med pipe-separatorer.
Bildrubriker (captions) placeras kursivt under bilden.
"""

import sys
import json
import zipfile
import os
import re
import shutil
from pathlib import Path

try:
    import docx
    import docx.text.paragraph
    import docx.table
except ImportError:
    print("ERROR: python-docx krävs. Kör: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Rubrikmappning: Word-rubrikstil → TKB-sektionsnamn
# ---------------------------------------------------------------------------
SECTION_SLUGS = {
    "1": "1-inledning",
    "2": "2-versionsinformation",
    "3": "3-tjanstedomanens-arkitektur",
    "4": "4-tjanstedomanens-krav-och-regler",
    "5": "5-tjanstedomanens-meddelandemodeller",
    "6": "6-gemensamma-informationskomponenter",
    "7": "7-tjanstekontrakt",
}

SECTION_TITLES = {
    "1": "1 Inledning",
    "2": "2 Versionsinformation",
    "3": "3 Tjänstedomänens arkitektur",
    "4": "4 Tjänstedomänens krav och regler",
    "5": "5 Tjänstedomänens meddelandemodeller",
    "6": "6 Gemensamma informationskomponenter",
    "7": "7 Tjänstekontrakt",
}


def heading_level(para) -> int:
    """Returnera rubriknivå (1-9) eller 0 om inte rubrik."""
    style = para.style.name if para.style else ""
    # Standard Word: "Heading 1", "Heading 2", etc.
    # Inera/svenska dokument: "Numrerad rubrik 1", "Rubrik 1", etc.
    for prefix in (
        "Heading ",
        "Rubrik ",
        "Numrerad rubrik ",
        "heading ",
        "rubrik ",
        "numrerad rubrik ",
    ):
        if style.lower().startswith(prefix.lower()):
            try:
                return int(style.split()[-1])
            except ValueError:
                pass
    # Numrerade listor som ser ut som rubriker — ignorera
    return 0


def get_image_rids(para) -> list:
    """Hämta relationship-IDs för inbäddade bilder i ett stycke."""
    rids = []
    blip_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    rel_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    for blip in para._p.findall(f".//{{{blip_ns}}}blip"):
        rid = blip.get(f"{{{rel_ns}}}embed")
        if rid:
            rids.append(rid)
    return rids


def extract_images(docx_path: Path, img_dir: Path) -> dict:
    """
    Extrahera alla mediafiler från docx till img_dir.
    Returnerar dict: rId → filnamn (relativt img_dir).
    """
    img_dir.mkdir(parents=True, exist_ok=True)
    doc = docx.Document(str(docx_path))
    rels = doc.part.rels
    rId_to_file = {}
    counter = 0

    with zipfile.ZipFile(str(docx_path)) as zf:
        zip_names = set(zf.namelist())

    for rId, rel in rels.items():
        if "image" not in rel.reltype.lower():
            continue
        try:
            part = rel.target_part
            partname = str(part.partname)  # e.g. /word/media/image1.png
            ext = partname.rsplit(".", 1)[-1].lower()
            if ext == "svg":
                # SVG: kopiera men konvertera inte — browsers stöder det
                pass
            counter += 1
            fname = f"img_{counter:03d}.{ext}"
            dest = img_dir / fname
            with open(str(dest), "wb") as f:
                f.write(part.blob)
            rId_to_file[rId] = fname
        except Exception as e:
            print(f"  WARN: kunde inte extrahera bild {rId}: {e}", file=sys.stderr)

    return rId_to_file


def table_to_md(table) -> str:
    """Konvertera en python-docx Table till Markdown-tabell."""
    if not table.rows:
        return ""

    # Samla celltexter — hantera sammanslagna celler
    rows_data = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            # Alla stycken i cellen, separerade med <br>-liknande mellanrum
            lines = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
            text = " / ".join(lines) if lines else ""
            # Escapa pipe-tecken
            text = text.replace("|", "\\|")
            cells.append(text)
        rows_data.append(cells)

    if not rows_data:
        return ""

    ncols = max(len(r) for r in rows_data)

    def pad_row(row):
        padded = row + [""] * (ncols - len(row))
        return padded[:ncols]

    lines = []
    header = pad_row(rows_data[0])
    lines.append("| " + " | ".join(header) + " |")
    lines.append("|" + " :--- |" * ncols)
    for row in rows_data[1:]:
        lines.append("| " + " | ".join(pad_row(row)) + " |")

    return "\n".join(lines)


def walk_body(body, doc_obj):
    """
    Iterera body-elementen i dokumentordning och yield:
        ('para', Paragraph)
        ('table', Table)
    Hanterar också <sdt> (structured document tags).
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def _walk(node):
        for child in node:
            local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if local == "p":
                yield ("para", docx.text.paragraph.Paragraph(child, doc_obj))
            elif local == "tbl":
                yield ("table", docx.table.Table(child, doc_obj))
            elif local in ("sdt", "body"):
                # structured doc tag eller body-nod — recurse
                sdt_content = child.find(f"{{{W}}}sdtContent")
                if sdt_content is not None:
                    yield from _walk(sdt_content)
                else:
                    yield from _walk(child)

    yield from _walk(body)


def detect_caption(para) -> bool:
    """
    Returnera True om stycket ser ut som en bildtext/figurbeskrivning.
    Vanliga stilar: Caption, Figur, Figure, Bildtext.
    """
    style = para.style.name if para.style else ""
    text = para.text.strip()
    style_lower = style.lower()
    text_lower = text.lower()
    return (
        "caption" in style_lower
        or "figur" in style_lower
        or "bildtext" in style_lower
        or text_lower.startswith("figur ")
        or text_lower.startswith("bild ")
        or text_lower.startswith("figure ")
        or (len(text) < 200 and re.match(r"^(figur|bild|figure|fig\.?)\s+\d+", text_lower))
    )


def build_full_markdown(docx_path: Path, img_dir: Path, img_rel_prefix: str = "images") -> tuple:
    """
    Bygg komplett Markdown för dokumentet.

    Returnerar:
        (markdown_str, structure_list)
        structure_list: [{"level": 1, "title": "Inledning", "pos": char_offset}, ...]
    """
    doc = docx.Document(str(docx_path))
    rId_to_file = extract_images(docx_path, img_dir)

    structure = []   # rubrikhierarki
    blocks = []      # markdown-block (strängar)
    img_used = set()

    prev_was_image = False

    for typ, elem in walk_body(doc.element.body, doc):
        if typ == "para":
            para = elem
            lvl = heading_level(para)
            text = para.text.strip()
            rids = get_image_rids(para)
            is_caption = detect_caption(para)

            if rids:
                for rid in rids:
                    if rid in rId_to_file and rid not in img_used:
                        img_used.add(rid)
                        fname = rId_to_file[rid]
                        alt = text if text else fname
                        blocks.append(f"\n![{alt}]({img_rel_prefix}/{fname})\n")
                prev_was_image = True
                # Om bilden har inline-text (sällsynt), lägg till
                if text and not is_caption:
                    blocks.append(text + "\n")

            elif (is_caption or (para.style and para.style.name == "Caption")) and prev_was_image:
                # Bildtext under bilden
                blocks.append(f"*{text}*\n")
                prev_was_image = False

            elif lvl > 0 and text:
                # Rubrik — vi lägger till en extra nivå (H1 i Word → ## i MD)
                # för att dokumenttiteln ska vara H1
                md_lvl = lvl + 1
                hashes = "#" * md_lvl
                blocks.append(f"\n{hashes} {text}\n")
                pos = sum(len(b) for b in blocks)
                structure.append({"level": lvl, "title": text, "md_level": md_lvl, "pos": pos})
                prev_was_image = False

            elif text:
                blocks.append(text + "\n")
                prev_was_image = False
            else:
                prev_was_image = False

        elif typ == "table":
            md_table = table_to_md(elem)
            if md_table:
                blocks.append(f"\n{md_table}\n")
            prev_was_image = False

    return "".join(blocks), structure


def split_into_sections(full_md: str, structure: list) -> dict:
    """
    Dela upp full_md i sektioner baserat på H2-rubriker (nivå 1 i Word).
    Returnerar dict: sektionsnummer (str) → markdown-text för sektionen.

    Matchar rubriktext mot kända sektionsnamn — Word-numrering är en listnumrering
    som inte ingår i den faktiska texten.
    """
    # Kändaa TKB-rubriker (titel → nummer)
    # Matchar på nyckelord som är unika nog per sektion
    TITLE_TO_NUM = {
        "inledning": "1",
        "versionsinformation": "2",
        "tjänstedomänens arkitektur": "3",
        "krav och regler": "4",
        "meddelandemodeller": "5",
        "gemensamma informationskomponenter": "6",
        "tjänstekontrakt": "7",
    }

    sections = {}
    # Hitta alla H2-rubriker i markdown — nivå 1 i Word → ## i output
    pattern = re.compile(r'^## (.+)$', re.MULTILINE)
    matches = list(pattern.finditer(full_md))

    for i, m in enumerate(matches):
        title = m.group(1).strip()
        title_lower = title.lower()

        # Försök matcha mot sektionsnyckelord
        num = None

        # Alternativ 1: Titeln börjar med siffra (t.ex. "1 Inledning")
        num_match = re.match(r'^(\d+)[\s\.]', title)
        if num_match:
            num = num_match.group(1)
        else:
            # Alternativ 2: Nyckelordsmatchning
            for keyword, section_num in TITLE_TO_NUM.items():
                if keyword in title_lower:
                    num = section_num
                    break

        if not num:
            continue

        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(full_md)
        sections[num] = full_md[start:end]

    return sections


def write_outputs(docx_path: Path, output_dir: Path):
    """Huvudfunktion — kör hela konverteringspipelinen."""
    output_dir.mkdir(parents=True, exist_ok=True)
    img_dir = output_dir / "images"
    sections_dir = output_dir / "sections"
    sections_dir.mkdir(exist_ok=True)

    print(f"Konverterar: {docx_path.name}")
    print(f"Output: {output_dir}")

    full_md, structure = build_full_markdown(docx_path, img_dir)

    # Spara hela dokumentet
    full_path = output_dir / "full-document.md"
    full_path.write_text(full_md, encoding="utf-8")
    print(f"  full-document.md: {len(full_md):,} tecken")

    # Spara rubrikkstruktur
    structure_path = output_dir / "structure.json"
    structure_path.write_text(json.dumps(structure, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"  structure.json: {len(structure)} rubriker")

    # Dela upp i sektioner
    sections = split_into_sections(full_md, structure)
    for num, slug in SECTION_SLUGS.items():
        fname = sections_dir / f"{slug}.md"
        if num in sections:
            content = sections[num]
            fname.write_text(content, encoding="utf-8")
            print(f"  sections/{slug}.md: {len(content):,} tecken")
        else:
            # Skapa tom sektion med notis
            fname.write_text(
                f"## {SECTION_TITLES.get(num, f'Avsnitt {num}')}\n\n"
                f"> **OBS:** Avsnitt {num} hittades inte i källdokumentet.\n",
                encoding="utf-8"
            )
            print(f"  sections/{slug}.md: SAKNAS i källdokumentet")

    # Bildstatistik
    img_files = list(img_dir.glob("*")) if img_dir.exists() else []
    print(f"  images/: {len(img_files)} filer ({', '.join(sorted(set(f.suffix for f in img_files)))})")

    print("Klar.")
    return {"sections": list(sections.keys()), "images": len(img_files), "structure": structure}


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(f"Användning: {sys.argv[0]} <docx-fil> <output-katalog>", file=sys.stderr)
        sys.exit(1)

    docx_path = Path(sys.argv[1])
    output_dir = Path(sys.argv[2])

    if not docx_path.exists():
        print(f"ERROR: filen {docx_path} finns inte.", file=sys.stderr)
        sys.exit(1)

    result = write_outputs(docx_path, output_dir)
    print(json.dumps(result, ensure_ascii=False, indent=2))
